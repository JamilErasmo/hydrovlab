import pdfplumber
from docx import Document

def pdf_to_word(pdf_path, docx_path):
    # Crear un nuevo documento Word
    doc = Document()

    # Abrir el archivo PDF
    with pdfplumber.open(pdf_path) as pdf:
        # Extraer texto de cada página
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()

            if text:
                # Añadir una nueva página en el documento Word
                doc.add_paragraph(f"Página {page_num + 1}")
                doc.add_paragraph(text)
                doc.add_page_break()

    # Guardar el documento Word
    doc.save(docx_path)
    print(f"Conversión completada. El archivo ha sido guardado como {docx_path}")

# Rutas de los archivos
pdf_path = "Correo_ ERICK LEONARDO CUENCA CAÑAR - Outlook.pdf"
docx_path = "archivo.docx"

# Llamar a la función para convertir PDF a Word
pdf_to_word(pdf_path, docx_path)
